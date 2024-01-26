import random
from docx import Document
from docx.shared import Pt  # 导入Pt类用于设置字体大小
from datetime import datetime, timedelta

# 定义生成四则运算题目的函数
def generate_math_question():
    operators = ['+', '-', '*', '/']  # 四种运算符
    operator = random.choice(operators)
    num1 = random.randint(1, 10000)  # 随机数范围根据需求调整
    if operator == '/' and num1 == 0:  # 确保除数不为零
        return generate_math_question()
    elif operator == '/':  # 对除法题目单独处理
        num2 = random.randint(1, num1 // 10)  # 确保有余数
        question = f"{num1} ÷ {num2}"
        answer = f"{num1 // num2} …… {num1 % num2}"
    elif operator == '*':
        num2 = random.randint(1, num1 - 1)
        question = f"{num1} X {num2}"
        answer = eval(f"{num1} * {num2}")
    else:
        num2 = random.randint(1, num1 - 1)
        question = f"{num1} {operator} {num2}"
        answer = eval(question)

    return question, answer

# 获取当前日期
current_date = datetime.now()

# 创建一个新的Word文档
doc = Document()

# 获取或新建一种样式，并修改其字体和字号
style = doc.styles['Normal']  # 获取默认的Normal样式
font = style.font
font.name = '微软雅黑'
font.size = Pt(14)

# 遍历未来几天
question_days = 12
for i in range(question_days):
    future_date = current_date + timedelta(days=i+1)
    date_string = future_date.strftime('%m月%d日')  # 格式化日期为'YYYY-MM-DD'
    
    # 在文档中添加题目页面
    questions_paragraph = doc.add_paragraph(f"{date_string} 题目：")

    # 添加分页
    doc.add_page_break()

    # 在文档中添加答案页面
    answers_paragraph = doc.add_paragraph(f"{date_string} 答案：")

    if i < question_days - 1:
        # 添加分页
        doc.add_page_break()

    for i in range(5):  # 根据需求调整题目数量
        question, _ = generate_math_question()
        questions_paragraph.add_run(f"\n{i+1}. {question} = ")
        # 除了最后一次，其余要添加空行
        if i < 4:
            for j in range(3):
                questions_paragraph.add_run(f"\n")
        answers_paragraph.add_run(f"\n{i+1}. {question} = {_}")


# 保存Word文档
doc.save('数学四则运算练习题和答案.docx')
